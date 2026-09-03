from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor

ROWS = [
    {"category": "System", "method": "GET", "route": "/api/status", "access": "Public", "description": "Checks whether the Next.js API layer is active and responding."},
    {"category": "Auth", "method": "POST", "route": "/api/auth/logout", "access": "Public", "description": "Ends a logged-out client session and clears auth state."},
    {"category": "Student", "method": "GET", "route": "/api/subjects", "access": "Student", "description": "Returns available subjects with optional filters for board, grade, or user context."},
    {"category": "Student", "method": "POST", "route": "/api/sessions/end", "access": "Student", "description": "Stores session duration and page-view analytics after a learning session ends."},
    {"category": "Internal", "method": "GET", "route": "/api/cron/assignment-reminders", "access": "Internal", "description": "Cron-triggered reminder job that sends pending assignment alerts to students."},
    {"category": "Student", "method": "POST", "route": "/api/assignments/generate", "access": "Student", "description": "Generates a new assignment set for a subject, topic, and difficulty level."},
    {"category": "Student", "method": "POST", "route": "/api/assignments/submit", "access": "Student", "description": "Submits answers for an assignment and returns scoring feedback."},
    {"category": "Student", "method": "GET", "route": "/api/practice/history", "access": "Student", "description": "Retrieves a student’s prior practice attempts and pagination data."},
    {"category": "Student", "method": "GET", "route": "/api/practice/metrics", "access": "Student", "description": "Provides test scores, attempts, average performance, recent activity, and streak metrics."},
    {"category": "Student", "method": "GET", "route": "/api/practice/search", "access": "Student", "description": "Searches prior practice attempts using subject, complexity, and keywords."},
    {"category": "Student", "method": "GET", "route": "/api/practice/subjects", "access": "Student", "description": "Lists practice subjects available for the logged-in child or user."},
    {"category": "Student", "method": "POST", "route": "/api/practice/generate", "access": "Student", "description": "Creates a new practice test from a topic and selected complexity."},
    {"category": "Student", "method": "POST", "route": "/api/practice/submit", "access": "Student", "description": "Submits practice responses and returns AI-based evaluation and scoring."},
    {"category": "Student", "method": "POST", "route": "/api/podcasts/generate", "access": "Student/Teacher", "description": "Generates podcast audio or answer-narration content for a topic or question."},
    {"category": "Student", "method": "GET", "route": "/api/student/progress", "access": "Student", "description": "Returns student progress trends such as tests taken, score averages, and streaks."},
    {"category": "Student", "method": "GET", "route": "/api/student/research-history", "access": "Student", "description": "Fetches the student’s AI research queries and prior exploration history."},
    {"category": "Student", "method": "GET", "route": "/api/student/preferences", "access": "Student", "description": "Loads dashboard theme and personalization preferences for the student profile."},
    {"category": "Student", "method": "POST", "route": "/api/student/preferences", "access": "Student", "description": "Saves customizable student settings such as theme, avatar, and gamification toggles."},
    {"category": "Student", "method": "GET", "route": "/api/student/notifications", "access": "Student", "description": "Retrieves the student notification feed and unread alert records."},
    {"category": "Student", "method": "POST", "route": "/api/student/notifications/read", "access": "Student", "description": "Marks selected student notifications as read."},
    {"category": "Student", "method": "GET", "route": "/api/student/gamification/stats", "access": "Student", "description": "Returns XP totals, streaks, and gamification status for the learner."},
    {"category": "Student", "method": "GET", "route": "/api/student/gamification/badges", "access": "Student", "description": "Fetches the badges awarded to the student for learning milestones and achievements."},
    {"category": "Student", "method": "GET", "route": "/api/student/gamification/leaderboard", "access": "Student", "description": "Shows leaderboard ranks and points for a student’s class or cohort."},
    {"category": "Student", "method": "GET", "route": "/api/student/teacher-assignments", "access": "Student", "description": "Lists teacher-assigned tasks assigned to a child and their current status."},
    {"category": "Student", "method": "POST", "route": "/api/student/teacher-assignments/[id]/submit", "access": "Student", "description": "Submits a teacher-assigned question response and stores evaluation data."},
    {"category": "Parent", "method": "GET", "route": "/api/parent/preferences", "access": "Parent", "description": "Loads parent-side settings, including gamification preferences for the child."},
    {"category": "Parent", "method": "POST", "route": "/api/parent/preferences", "access": "Parent", "description": "Updates parent-managed preferences for the child account or portal."},
    {"category": "Parent", "method": "GET", "route": "/api/parent/notifications", "access": "Parent", "description": "Retrieves notifications relevant to the parent account."},
    {"category": "Parent", "method": "POST", "route": "/api/parent/notifications/read", "access": "Parent", "description": "Marks a parent’s notifications as read after they are reviewed."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/auth/register", "access": "Public", "description": "Registers a teacher account, creates credentials, and sends verification email."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/auth/verify-email", "access": "Public", "description": "Verifies the teacher email using a registration token."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/auth/resend-verification", "access": "Public", "description": "Resends the verification email if the teacher did not receive it."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/auth/login", "access": "Public", "description": "Authenticates a teacher using credentials and returns their personal session data."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/auth/logout", "access": "Teacher", "description": "Logs out the current teacher session and invalidates client cookies."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/analytics", "access": "Teacher", "description": "Provides class performance, student totals, and summary-ledger metrics."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/students", "access": "Teacher", "description": "Lists all students associated with the teacher’s managed classes."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/students/search", "access": "Teacher", "description": "Searches students by name or metadata for quick classroom management."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/classes", "access": "Teacher", "description": "Returns teacher-owned or assigned classes and class metadata."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/classes", "access": "Teacher", "description": "Creates a new class with board, grade, subjects, and naming data."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/classes/[classId]/invite-token", "access": "Teacher", "description": "Generates an invite token and link for students to join a class."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/classes/[classId]/enroll", "access": "Teacher", "description": "Enrolls a child or invited learner into a class using token or child ID."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/classes/[classId]/students", "access": "Teacher", "description": "Loads student roster and class activity for a specific class."},
    {"category": "Teacher", "method": "DELETE", "route": "/api/teacher/classes/[classId]/students", "access": "Teacher", "description": "Removes a student from the class roster when required."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/assignments", "access": "Teacher", "description": "Lists teacher-generated assignments across classes and due dates."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/assignments/generate", "access": "Teacher", "description": "Builds a classroom assignment with topic, complexity, due date, and question payload."},
    {"category": "Teacher", "method": "PATCH", "route": "/api/teacher/assignments/[id]", "access": "Teacher", "description": "Updates assignment metadata, including due date and learner targeting."},
    {"category": "Teacher", "method": "PUT", "route": "/api/teacher/assignments/[id]", "access": "Teacher", "description": "Handles assignment update or status operations for a teacher-managed task."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/assignments/[id]/status", "access": "Teacher", "description": "Shows submission progress and overall performance for a given assignment."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/assignments/[id]/review", "access": "Teacher", "description": "Fetches assignment submissions for grading and teacher review."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/assignments/[id]/review", "access": "Teacher", "description": "Submits a score and feedback note for a student’s assignment submission."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/assignments/[id]/remind", "access": "Teacher", "description": "Sends reminders to selected students with pending assignment tasks."},
    {"category": "Teacher", "method": "GET", "route": "/api/teacher/notifications", "access": "Teacher", "description": "Returns teacher notifications and important platform updates."},
    {"category": "Teacher", "method": "POST", "route": "/api/teacher/notifications/read", "access": "Teacher", "description": "Marks teacher notifications as read after they are acted upon."},
    {"category": "Admin", "method": "POST", "route": "/api/admin/auth/login", "access": "Public", "description": "Authenticates a system administrator for privileged dashboard access."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/verify-session", "access": "Admin", "description": "Checks whether the current admin user is still authenticated and authorized."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/accounts", "access": "Admin", "description": "Lists administrator accounts and account-level access metadata."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/users", "access": "Admin", "description": "Returns a paginated list of users and search results for admin moderation."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/users/[id]", "access": "Admin", "description": "Fetches a single user record with profile and subject information."},
    {"category": "Admin", "method": "POST", "route": "/api/admin/invite", "access": "Admin", "description": "Creates an invite token for onboarding a new admin or support team member."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/accept-invite", "access": "Public", "description": "Validates a pending admin invitation token to confirm account details."},
    {"category": "Admin", "method": "POST", "route": "/api/admin/accept-invite", "access": "Public", "description": "Completes admin account setup using the invitation token and password."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/credits", "access": "Admin", "description": "Displays AI/feature usage logs, token consumption, and cost summaries."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/features/podcast", "access": "Admin", "description": "Lists student and teacher podcast access records and feature states."},
    {"category": "Admin", "method": "POST", "route": "/api/admin/features/podcast", "access": "Admin", "description": "Enables or disables podcast access for a user role and saves admin notes."},
    {"category": "Admin", "method": "GET", "route": "/api/admin/notifications", "access": "Admin", "description": "Retrieves admin notifications for operational and platform alerts."},
    {"category": "Admin", "method": "POST", "route": "/api/admin/notifications/read", "access": "Admin", "description": "Marks the selected admin notifications as read."},
]


def shade_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    shd.set(qn('w:color'), 'auto')


def style_header_row(table):
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
            if not paragraph.runs:
                p_run = paragraph.add_run()
                p_run.bold = True
                p_run.font.name = 'Calibri'
                p_run.font.size = Pt(10)
                p_run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cell, '2F75B5')
        cell.vertical_alignment = None


def add_highlighted_line(doc, text, color='yellow'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.highlight_color = getattr(WD_COLOR_INDEX, color.upper(), WD_COLOR_INDEX.YELLOW)


def build_document(output_path: str):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('School Project API Inventory')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Application APIs and route summaries')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.name = 'Calibri'

    add_highlighted_line(doc, 'Key highlight: Core authentication, student learning flows, teacher classroom management, parent notifications, and admin controls.')
    doc.add_paragraph('Prepared from the application’s API route definitions and project documentation. This catalog reflects the main routes currently developed for the platform.')

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = 'Category'
    table.rows[0].cells[1].text = 'Method'
    table.rows[0].cells[2].text = 'Route'
    table.rows[0].cells[3].text = 'Access'
    table.rows[0].cells[4].text = 'Brief Description'
    style_header_row(table)

    for row in ROWS:
        cells = table.add_row().cells
        cells[0].text = row['category']
        cells[1].text = row['method']
        cells[2].text = row['route']
        cells[3].text = row['access']
        cells[4].text = row['description']

    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

    doc.add_paragraph('')
    doc.add_paragraph('Notes:')
    notes = doc.add_paragraph()
    notes_run = notes.add_run('• Most endpoints use JSON payloads and return standard error responses, including 400, 401, 403, and 500 status codes.\n• Internal cron routes are protected with authorization headers for automated background jobs.\n• Some routes are public for onboarding or verification while others are restricted to Student, Teacher, Parent, or Admin access.')
    notes_run.font.name = 'Calibri'
    notes_run.font.size = Pt(10)

    doc.save(output_path)
    print(f'Created Word document: {output_path}')


if __name__ == '__main__':
    build_document(r'd:\Shyam\School Project\School Project Workspace\docs\School_Project_API_Document.docx')
